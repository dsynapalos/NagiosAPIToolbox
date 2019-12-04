import os
import re
from datetime import datetime, timedelta
import docx
import csv
from docx import Document
from docx.enum.dml import *
from docx.enum.table import *
from docx.enum.text import *
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Cm, RGBColor
import pandas as pd
import numpy as np
import main as m
import matplotlib.pyplot as plt
import win32com.client as client
from docx.enum.table import *


#Edited from https://stackoverflow.com/questions/32992457/update-the-toc-table-of-content-of-ms-word-docx-documents-with-python
def update_toc(docx_file):
    word = client.DispatchEx("Word.Application")
    try:
        doc = word.Documents.Open(docx_file)
        doc.TablesOfContents(1).Update()
        doc.Close(SaveChanges=True)
    finally:
        word.Quit()
        del word

#Edited from https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#Edited from https://stackoverflow.com/questions/55572685/how-set-direction-of-table-of-content-in-docx-files-uisng-python-docx
def create_Toc(cr_document):

    cont = cr_document.add_heading('Table of Contents', 8)
    cont.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = cr_document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'  # change 1-4 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    cr_document.add_page_break()


def line_brake(doc, n):
    for i in range(0, n):
        doc.add_paragraph()

def change_orientation(document):
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width

    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

def csv2chart(file,doc):
    with open(file, newline='') as csv_file:
        csv_reader = csv.reader(csv_file)
        csv_headers = next(csv_reader)

        csv_cols = len(csv_headers)

        table = doc.add_table(rows=1, cols=csv_cols)
        table.style = 'LightGrid-Accent1'
        table.autofit = False
        col = table.columns[0]
        col.width = Cm(1)
        hdr_cells = table.rows[0].cells

        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(csv_cols):
                row_cells[i].text = row[i]



#https://github.com/python-openxml/python-docx/issues/113
def convert_to_pdf(filepath:str):
    """Save a pdf of a docx file."""
    try:
        word = client.DispatchEx("Word.Application")
        target_path = filepath.replace(".docx", r".pdf")
        word_doc = word.Documents.Open(filepath)
        word_doc.SaveAs(target_path, FileFormat=17)
        word_doc.Close()
    except Exception as e:
            raise e
    finally:
            word.Quit()
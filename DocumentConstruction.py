from docx import Document, opc
from docx.enum.text import *
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import *
from datetime import datetime
import clients as CL
import main as MN
import DOCXFunctions as DF


def constructDocument(index, template):
    new_document = Document(template)

    '''
    Preface
    '''

    new_document.add_paragraph().add_run().add_picture(CL.company_logo, width=Cm(12.7))
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 1)

    new_document.add_paragraph().add_run(CL.report_type).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 2)

    new_document.add_paragraph().add_run(CL.Clients[index]['TITLE']).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 2)

    new_document.add_paragraph().add_run(str(datetime.now().strftime('%d/%m/%Y'))).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 2)

    new_document.add_paragraph().add_run(CL.report_desc).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 6)

    new_document.add_paragraph().add_run(CL.classify).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    new_document.add_page_break()

    '''
    Contacts
    '''

    new_document.add_paragraph().add_run(CL.inform).font.name = 'Cambria'
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 3)

    new_document.add_paragraph().add_run().add_picture(CL.Clients[index]['INFO_CHART'], width=Cm(15.5))
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    new_document.add_page_break()

    '''
    Table Of Contents
    '''

    paragraph = new_document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-4 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    new_document.add_page_break()

    '''
    Chapter 1
    '''

    new_document.add_heading(CL.chapter1_0, 1)

    DF.line_brake(new_document, 1)

    new_document.add_heading(CL.chapter1_1, 2)

    new_document.add_page_break()

    new_document.add_heading(CL.chapter1_2, 2)

    new_document.add_page_break()

    new_document.add_heading(CL.chapter1_3, 2)

    new_document.add_page_break()

    DF.change_orientation(new_document)

    new_document.add_heading(CL.chapter1_4, 2)

    for host in CL.Clients[CL.Index]['BAND_GRAPHS'].keys():
        for number in range(0, len(CL.Clients[CL.Index]['BAND_GRAPHS'][host])):
            new_document.add_picture(
                MN.DIRECTORY + '/Screenshots/' + host + '_' + CL.Clients[CL.Index]['BAND_GRAPHS'][host][number].replace(
                    ' ', '').replace('-', '_') + '.png', width=Cm(22.25))

    new_document.add_page_break()

    '''
    Chapter 2
    '''
    DF.change_orientation(new_document)

    new_document.add_heading(CL.chapter2_0, 1)

    DF.line_brake(new_document, 2)

    new_document.add_paragraph().add_run(CL.chapter2_0_1)

    DF.line_brake(new_document, 1)

    new_document.add_paragraph().add_run(CL.chapter2_0_2)
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    DF.line_brake(new_document, 1)

    new_document.add_heading(CL.chapter2_1, 2)

    DF.line_brake(new_document, 1)

    DF.csv2chart(MN.DIRECTORY + '/Stats Working Dir/Important Bandwidth.csv', new_document)

    new_document.add_page_break()

    new_document.add_heading(CL.chapter2_2, 2)

    DF.line_brake(new_document, 1)

    DF.csv2chart(MN.DIRECTORY + '/Stats Working Dir/Internet Bandwidth.csv', new_document)

    new_document.add_page_break()

    '''
    Chapter 3
    '''

    new_document.add_heading(CL.chapter3_0, 1)

    DF.line_brake(new_document, 2)

    new_document.add_paragraph().add_run(CL.chapter3_0_1)

    DF.line_brake(new_document, 1)

    new_document.add_paragraph().add_run(CL.chapter3_0_2)
    new_document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    new_document.add_heading(CL.chapter3_1, 2)

    DF.line_brake(new_document, 1)

    DF.csv2chart(MN.DIRECTORY + '/Stats Working Dir/Host Downtimes.csv', new_document)

    new_document.add_page_break()

    new_document.add_heading(CL.chapter3_2, 2)

    DF.line_brake(new_document, 1)

    DF.csv2chart(MN.DIRECTORY + '/Stats Working Dir/Interface Downtimes.csv', new_document)

    #new_document.add_page_break()

    new_document.save(MN.DIRECTORY + '/demo.docx')

    DF.update_toc(MN.DIRECTORY + '/demo.docx')

    DF.convert_to_pdf(MN.DIRECTORY + '/demo.docx')
